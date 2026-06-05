import asyncio
import io
import logging
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
import pypdf
from docx import Document
from PIL import Image

try:
    import pytesseract
except ImportError:
    pytesseract = None

logger = logging.getLogger(__name__)

EXTRACTION_TIMEOUT = 30  # seconds
EXTRACTION_TIMEOUT_DOC = 60  # DOC conversion via libreoffice is slower


class ExtractorError(Exception):
    """Base extractor error."""
    pass


class ExtractorTimeoutError(ExtractorError):
    """Extraction timed out."""
    pass


class ExtractorFormatError(ExtractorError):
    """Unsupported file format."""
    pass


def detect_file_format(content: bytes, filename: Optional[str] = None) -> str:
    """Detect file format from magic bytes, with filename as a tie-breaker.

    Supported formats: pdf, docx, doc, rtf, odt, txt
    """
    if not content:
        return "unknown"

    # PDF: %PDF-
    if content[:4] == b"%PDF":
        return "pdf"

    # Images (handled via ConvertAPI OCR — local stack can't read these):
    if content[:3] == b"\xff\xd8\xff":            # JPEG
        return "jpg"
    if content[:8] == b"\x89PNG\r\n\x1a\n":       # PNG
        return "png"
    if content[:4] in (b"II*\x00", b"MM\x00*"):   # TIFF (LE / BE)
        return "tiff"
    if content[:4] == b"RIFF" and content[8:12] == b"WEBP":  # WEBP
        return "webp"

    # RTF: {\rtf
    if content[:5] == b"{\\rtf":
        return "rtf"

    # MS Compound Document (DOC, XLS, PPT - Office 97-2003): D0 CF 11 E0 A1 B1 1A E1
    if content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        # Could be .doc/.xls/.ppt - use filename hint if available
        if filename:
            fl = filename.lower()
            if fl.endswith(".doc"):
                return "doc"
            if fl.endswith((".xls", ".xlsx", ".ppt", ".pptx")):
                return "unknown"  # not a CV format
        # Default to DOC since that's what we mostly see in CV uploads
        return "doc"

    # ZIP-based formats: PK\x03\x04 or PK\x05\x06 (empty) or PK\x07\x08 (spanned)
    if content[:2] == b"PK":
        # Probe content for OOXML markers (search first 64KB only for performance)
        probe = content[:65536]
        if b"word/document.xml" in probe or b"word/_rels" in probe:
            return "docx"
        if b"mimetypeapplication/vnd.oasis.opendocument.text" in probe:
            return "odt"
        if b"xl/workbook.xml" in probe or b"ppt/presentation.xml" in probe:
            return "unknown"  # spreadsheet / presentation
        # Trust filename extension as last resort for ZIP archives
        if filename:
            fl = filename.lower()
            if fl.endswith(".docx"):
                return "docx"
            if fl.endswith(".odt"):
                return "odt"
        # Default to DOCX since that's the most common for CV uploads
        return "docx"

    # Plain text heuristic: mostly printable ASCII / UTF-8
    if filename and filename.lower().endswith(".txt"):
        return "txt"
    try:
        sample = content[:4096].decode("utf-8")
        printable_ratio = sum(1 for c in sample if c.isprintable() or c in "\r\n\t") / max(len(sample), 1)
        if printable_ratio > 0.95:
            return "txt"
    except UnicodeDecodeError:
        pass

    return "unknown"


async def extract_text_from_pdf(content: bytes) -> tuple[str, str]:
    """
    Extract text from PDF using pypdf, fallback to PyMuPDF.

    Returns:
        Tuple of (extracted_text, method_used)
    """

    async def _pypdf_extract() -> tuple[str, str]:
        """Extract using pypdf."""
        try:
            pdf_file = io.BytesIO(content)
            reader = pypdf.PdfReader(pdf_file)
            text_parts = []

            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")

            text = "\n".join(text_parts)
            if not text.strip():
                raise ValueError("No text extracted from PDF")

            logger.debug(f"Extracted {len(text)} chars from PDF using pypdf")
            return text, "pypdf"

        except Exception as e:
            logger.debug(f"pypdf extraction failed: {e}")
            raise

    async def _pymupdf_extract() -> tuple[str, str]:
        """Extract using PyMuPDF (fallback)."""
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page_num, page in enumerate(doc):
                try:
                    text = page.get_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num} with PyMuPDF: {e}")

            text = "\n".join(text_parts)
            doc.close()

            if not text.strip():
                raise ValueError("No text extracted from PDF")

            logger.debug(f"Extracted {len(text)} chars from PDF using PyMuPDF")
            return text, "pymupdf"

        except Exception as e:
            logger.debug(f"PyMuPDF extraction failed: {e}")
            raise

    # Try pypdf first
    try:
        return await asyncio.wait_for(_pypdf_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("PDF extraction timed out with pypdf")
    except Exception as e:
        logger.debug(f"pypdf failed, trying PyMuPDF: {e}")

    # Fallback to PyMuPDF
    try:
        return await asyncio.wait_for(_pymupdf_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("PDF extraction timed out with PyMuPDF")
    except Exception as e:
        logger.error(f"All PDF extraction methods failed: {e}")
        raise ExtractorFormatError(f"Failed to extract PDF: {e}")


async def extract_text_from_docx(content: bytes) -> tuple[str, str]:
    """
    Extract text from DOCX using python-docx.

    Returns:
        Tuple of (extracted_text, method_used)
    """

    async def _docx_extract() -> tuple[str, str]:
        """Extract using python-docx."""
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            text = "\n".join(text_parts)
            if not text.strip():
                raise ValueError("No text extracted from DOCX")

            logger.debug(f"Extracted {len(text)} chars from DOCX")
            return text, "docx"

        except Exception as e:
            logger.debug(f"DOCX extraction failed: {e}")
            raise

    try:
        return await asyncio.wait_for(_docx_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("DOCX extraction timed out")
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ExtractorFormatError(f"Failed to extract DOCX: {e}")


async def extract_text_with_ocr(content: bytes) -> tuple[str, str]:
    """
    Extract text from PDF using OCR (pytesseract).

    Returns:
        Tuple of (extracted_text, method_used)
    """
    if pytesseract is None:
        raise ExtractorError("pytesseract not available, cannot perform OCR")

    async def _ocr_extract() -> tuple[str, str]:
        """Extract using OCR."""
        try:
            # Convert PDF to images
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page_num, page in enumerate(doc):
                try:
                    # Render page to image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    # OCR
                    text = pytesseract.image_to_string(img, lang="eng+heb")
                    if text.strip():
                        text_parts.append(text)

                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")

            doc.close()
            text = "\n".join(text_parts)

            if not text.strip():
                raise ValueError("No text extracted via OCR")

            logger.debug(f"Extracted {len(text)} chars from PDF using OCR")
            return text, "ocr"

        except Exception as e:
            logger.debug(f"OCR extraction failed: {e}")
            raise

    try:
        return await asyncio.wait_for(_ocr_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("OCR extraction timed out")
    except Exception as e:
        logger.error(f"OCR extraction failed: {e}")
        raise ExtractorFormatError(f"Failed to extract via OCR: {e}")


async def extract_text_from_doc(content: bytes) -> tuple[str, str]:
    """
    Extract text from legacy DOC (Word 97-2003 binary) files.

    Strategy (in order):
      1. antiword (fastest, no GUI deps) - if installed
      2. catdoc (fast, common on linux) - if installed
      3. soffice / libreoffice headless (slowest, heaviest) - if installed
      4. Best-effort raw text strip (last resort)
    """

    async def _run_subprocess(cmd: list[str], input_bytes: Optional[bytes] = None,
                               work_dir: Optional[str] = None) -> Optional[str]:
        """Run a subprocess asynchronously, return stdout text or None on failure."""
        try:
            proc = await asyncio.create_subprocess_exec(
                *cmd,
                stdin=asyncio.subprocess.PIPE if input_bytes else None,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                cwd=work_dir,
            )
            stdout, stderr = await proc.communicate(input=input_bytes)
            if proc.returncode != 0:
                logger.debug(f"{cmd[0]} returned {proc.returncode}: {stderr.decode(errors='ignore')[:200]}")
                return None
            return stdout.decode("utf-8", errors="ignore")
        except FileNotFoundError:
            return None
        except Exception as e:
            logger.debug(f"Subprocess {cmd[0]} failed: {e}")
            return None

    async def _doc_extract() -> tuple[str, str]:
        # 1. antiword (read from stdin)
        if shutil.which("antiword"):
            text = await _run_subprocess(["antiword", "-"], input_bytes=content)
            if text and text.strip():
                logger.debug(f"Extracted {len(text)} chars from DOC using antiword")
                return text, "antiword"

        # 2. catdoc (read from stdin)
        if shutil.which("catdoc"):
            text = await _run_subprocess(["catdoc", "-"], input_bytes=content)
            if text and text.strip():
                logger.debug(f"Extracted {len(text)} chars from DOC using catdoc")
                return text, "catdoc"

        # 3. soffice / libreoffice (convert to txt in temp dir)
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            with tempfile.TemporaryDirectory() as tmpdir:
                doc_path = Path(tmpdir) / "input.doc"
                doc_path.write_bytes(content)
                conv = await _run_subprocess(
                    [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", tmpdir, str(doc_path)],
                    work_dir=tmpdir,
                )
                if conv is not None:
                    txt_path = Path(tmpdir) / "input.txt"
                    if txt_path.exists():
                        text = txt_path.read_text(encoding="utf-8", errors="ignore")
                        if text.strip():
                            logger.debug(f"Extracted {len(text)} chars from DOC using libreoffice")
                            return text, "libreoffice"

        # 4. Best-effort: extract printable ASCII/UTF-8 runs from the binary stream.
        # DOC binary contains plenty of structural junk, but readable text is
        # usually grouped in long runs. This salvages content when no converter
        # is installed.
        try:
            # Find runs of 4+ printable chars (ASCII + Hebrew + common Latin extended)
            text_bytes = re.findall(rb"[\x20-\x7e\xc0-\xff\xd0-\xd7\xd8-\xff\n\r\t]{4,}", content)
            decoded = b"\n".join(text_bytes).decode("utf-8", errors="ignore")
            # Also try cp1255 for Hebrew, cp1252 for Western European
            for encoding in ("cp1255", "cp1252", "latin-1"):
                try:
                    alt = b"\n".join(text_bytes).decode(encoding, errors="ignore")
                    if len(alt) > len(decoded):
                        decoded = alt
                except Exception:
                    pass
            if decoded.strip() and len(decoded) > 100:
                logger.warning(
                    f"DOC extracted via best-effort strip ({len(decoded)} chars) - "
                    "install antiword/catdoc/libreoffice for better quality"
                )
                return decoded, "best-effort-strip"
        except Exception as e:
            logger.debug(f"Best-effort DOC strip failed: {e}")

        raise ExtractorFormatError(
            "Cannot extract DOC: install antiword, catdoc, or libreoffice"
        )

    try:
        return await asyncio.wait_for(_doc_extract(), timeout=EXTRACTION_TIMEOUT_DOC)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("DOC extraction timed out")
    except ExtractorError:
        raise
    except Exception as e:
        logger.error(f"DOC extraction failed: {e}")
        raise ExtractorFormatError(f"Failed to extract DOC: {e}")


async def extract_text_from_rtf(content: bytes) -> tuple[str, str]:
    """Extract text from RTF using striprtf (pure-python) or a regex fallback."""

    async def _rtf_extract() -> tuple[str, str]:
        # Decode RTF bytes - RTF is 7-bit ASCII with escapes for unicode/cp1252
        try:
            rtf_str = content.decode("utf-8", errors="ignore")
        except Exception:
            rtf_str = content.decode("latin-1", errors="ignore")

        # Try striprtf if available (best quality)
        try:
            from striprtf.striprtf import rtf_to_text
            text = rtf_to_text(rtf_str, errors="ignore")
            if text and text.strip():
                logger.debug(f"Extracted {len(text)} chars from RTF using striprtf")
                return text, "striprtf"
        except ImportError:
            logger.debug("striprtf not installed, using regex fallback")

        # Regex fallback - strip RTF control words and groups
        # Remove control words: \word, \word123, \word-123
        text = re.sub(r"\\[a-zA-Z]+-?\d*\s?", " ", rtf_str)
        # Remove escaped chars: \\, \{, \}, \'hh
        text = re.sub(r"\\\\|\\\{|\\\}", "", text)
        text = re.sub(r"\\'[0-9a-fA-F]{2}", "", text)
        # Remove braces and pictures
        text = re.sub(r"\{\\\*[^}]*\}", "", text)
        text = re.sub(r"[{}]", "", text)
        # Collapse whitespace
        text = re.sub(r"\s+", " ", text).strip()

        if not text:
            raise ValueError("No text extracted from RTF")

        logger.debug(f"Extracted {len(text)} chars from RTF using regex fallback")
        return text, "rtf-regex"

    try:
        return await asyncio.wait_for(_rtf_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("RTF extraction timed out")
    except Exception as e:
        logger.error(f"RTF extraction failed: {e}")
        raise ExtractorFormatError(f"Failed to extract RTF: {e}")


async def extract_text_from_odt(content: bytes) -> tuple[str, str]:
    """Extract text from OpenDocument Text using zipfile + XML parsing."""

    async def _odt_extract() -> tuple[str, str]:
        import zipfile
        from xml.etree import ElementTree as ET

        try:
            with zipfile.ZipFile(io.BytesIO(content)) as z:
                if "content.xml" not in z.namelist():
                    raise ValueError("ODT missing content.xml")
                xml_bytes = z.read("content.xml")

            # ODT uses namespaced XML; strip namespaces for simple text extraction
            root = ET.fromstring(xml_bytes)
            text_parts = []
            for elem in root.iter():
                if elem.text and elem.text.strip():
                    text_parts.append(elem.text)
                if elem.tail and elem.tail.strip():
                    text_parts.append(elem.tail)

            text = "\n".join(text_parts)
            if not text.strip():
                raise ValueError("No text in ODT")

            logger.debug(f"Extracted {len(text)} chars from ODT")
            return text, "odt"
        except Exception as e:
            logger.debug(f"ODT extraction failed: {e}")
            raise

    try:
        return await asyncio.wait_for(_odt_extract(), timeout=EXTRACTION_TIMEOUT)
    except asyncio.TimeoutError:
        raise ExtractorTimeoutError("ODT extraction timed out")
    except Exception as e:
        logger.error(f"ODT extraction failed: {e}")
        raise ExtractorFormatError(f"Failed to extract ODT: {e}")


async def extract_text_from_txt(content: bytes) -> tuple[str, str]:
    """Decode plain text, trying multiple encodings for Hebrew support."""
    for encoding in ("utf-8", "utf-8-sig", "cp1255", "cp1252", "latin-1"):
        try:
            text = content.decode(encoding)
            if text.strip():
                logger.debug(f"Decoded {len(text)} chars from TXT as {encoding}")
                return text, f"txt-{encoding}"
        except UnicodeDecodeError:
            continue
    raise ExtractorFormatError("Could not decode plain text file")


# Public list of formats this module can handle. Useful for diagnostics.
SUPPORTED_FORMATS = ("pdf", "docx", "doc", "rtf", "odt", "txt")


_MIN_USEFUL_TEXT = 30  # chars; below this we treat extraction as a "miss"


async def _try_convertapi(filename: str, content: bytes) -> Optional[str]:
    """Attempt extraction via ConvertAPI. Returns text on success, else None.

    Never raises — on any error/misconfig returns None so the caller uses the
    local extractors. No-op when no secret is configured."""
    try:
        from pandapower.integrations.convertapi_client import (
            ConvertApiClient,
            convertapi_src_token,
            convertapi_within_budget,
            get_convertapi_config,
        )

        cfg = await get_convertapi_config()
        if not cfg.get("enabled") or not cfg.get("secret"):
            return None

        # BUDGET GUARD: stop using ConvertAPI before we exceed the plan limit
        # (overage is billed extra). When over budget we return None so the
        # caller falls back to the free local extractors — the pipeline keeps
        # moving, just without managed OCR for scanned/image CVs.
        allowed, usage = await convertapi_within_budget(cfg)
        if not allowed:
            consumed = (usage or {}).get("consumed")
            total = (usage or {}).get("total")
            logger.warning(
                f"[convertapi] BUDGET REACHED ({consumed}/{total}, "
                f"max={cfg.get('max_usage_pct')}) — skipping ConvertAPI for "
                f"{filename}, using local extractors to avoid overage charges."
            )
            return None

        file_format = detect_file_format(content, filename=filename)
        src = convertapi_src_token(file_format, filename)
        if not src:
            logger.info(f"ConvertAPI: no source mapping for {filename} (format={file_format})")
            return None

        client = ConvertApiClient(cfg["secret"])
        try:
            text = await client.to_text(content, src, ocr_languages=cfg.get("ocr_languages", "en,he"))
        finally:
            await client.close()

        if text and len(text.strip()) >= _MIN_USEFUL_TEXT:
            logger.info(f"ConvertAPI extracted {len(text)} chars from {filename} (src={src})")
            return text
        logger.info(f"ConvertAPI returned too little text for {filename} ({len(text or '')} chars)")
        return None
    except Exception as e:
        logger.warning(f"ConvertAPI extraction failed for {filename} (falling back to local): {e}")
        return None


async def extract_text(filename: str, content: bytes) -> tuple[str, str]:
    """Public orchestrator. Routes through ConvertAPI (managed OCR) per config,
    with the local extractor stack as the safety net.

    - mode "always":   ConvertAPI first; on miss/error → local extractors.
    - mode "fallback": local extractors first; on miss/error → ConvertAPI.
    - no secret:       behaves exactly like the local-only pipeline (no regression).
    """
    try:
        from pandapower.integrations.convertapi_client import get_convertapi_config
        cfg = await get_convertapi_config()
        mode = (cfg.get("mode") or "always").lower()
        enabled = bool(cfg.get("enabled") and cfg.get("secret"))
    except Exception:
        mode, enabled = "always", False

    if enabled and mode != "fallback":
        text = await _try_convertapi(filename, content)
        if text is not None:
            return text, "convertapi"
        return await _extract_text_local(filename, content)

    # fallback mode (or ConvertAPI disabled): try local first.
    try:
        text, method = await _extract_text_local(filename, content)
        if text and len(text.strip()) >= _MIN_USEFUL_TEXT:
            return text, method
        local_result = (text, method)
    except (ExtractorError, ExtractorTimeoutError):
        local_result = None

    if enabled:
        ca_text = await _try_convertapi(filename, content)
        if ca_text is not None:
            return ca_text, "convertapi"

    if local_result is not None:
        return local_result
    raise ExtractorFormatError(f"All extraction methods failed for {filename}")


async def _extract_text_local(filename: str, content: bytes) -> tuple[str, str]:
    """
    Local-only orchestrator: detect format and extract text with fallback routing.

    Supports: PDF, DOCX, DOC (Word 97-2003), RTF, ODT, TXT.

    For PDFs that yield no text (scanned images), automatically falls back to OCR
    if pytesseract is installed.

    Args:
        filename: Original filename (used as a tie-breaker for ambiguous formats)
        content: File content bytes

    Returns:
        Tuple of (extracted_text, method_used)

    Raises:
        ExtractorFormatError: If format is unsupported or extraction failed
        ExtractorTimeoutError: If extraction times out
    """
    logger.debug(f"Extracting text from {filename} ({len(content)} bytes)")

    file_format = detect_file_format(content, filename=filename)
    logger.info(f"File {filename}: detected format = {file_format}")

    if file_format == "pdf":
        # Try text-based extraction first
        try:
            text, method = await extract_text_from_pdf(content)
            # If the PDF returned almost no text, it's probably scanned - try OCR
            if len(text.strip()) < 50 and pytesseract is not None:
                logger.info(f"PDF {filename} yielded only {len(text)} chars - attempting OCR fallback")
                try:
                    ocr_text, ocr_method = await extract_text_with_ocr(content)
                    if len(ocr_text.strip()) > len(text.strip()):
                        return ocr_text, ocr_method
                except ExtractorError as ocr_err:
                    logger.debug(f"OCR fallback failed: {ocr_err}")
            return text, method
        except ExtractorTimeoutError:
            logger.warning(f"PDF text extraction timed out for {filename}")
            raise
        except ExtractorError as e:
            # Text extraction failed entirely - try OCR
            if pytesseract is not None:
                logger.info(f"PDF text extraction failed for {filename}, trying OCR: {e}")
                try:
                    return await extract_text_with_ocr(content)
                except ExtractorError as ocr_err:
                    logger.error(f"Both text and OCR extraction failed for {filename}: {ocr_err}")
                    raise ExtractorFormatError(f"PDF extraction failed (text + OCR): {ocr_err}")
            raise

    elif file_format == "docx":
        return await extract_text_from_docx(content)

    elif file_format == "doc":
        return await extract_text_from_doc(content)

    elif file_format == "rtf":
        return await extract_text_from_rtf(content)

    elif file_format == "odt":
        return await extract_text_from_odt(content)

    elif file_format == "txt":
        return await extract_text_from_txt(content)

    else:
        # Last-ditch: try by extension if magic-byte detection failed
        if filename:
            fl = filename.lower()
            if fl.endswith(".pdf"):
                return await extract_text_from_pdf(content)
            if fl.endswith(".docx"):
                return await extract_text_from_docx(content)
            if fl.endswith(".doc"):
                return await extract_text_from_doc(content)
            if fl.endswith(".rtf"):
                return await extract_text_from_rtf(content)
            if fl.endswith(".odt"):
                return await extract_text_from_odt(content)
            if fl.endswith(".txt"):
                return await extract_text_from_txt(content)

        raise ExtractorFormatError(
            f"Unsupported file format for {filename}: detected '{file_format}'. "
            f"Supported formats: {', '.join(SUPPORTED_FORMATS)}"
        )
